from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/strategy/market-opportunity-platform-transition-brief.md"
OUTPUT = ROOT / "docs/strategy/Market_Opportunity_Platform_Transition_Brief.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="C8D0DA", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, size=9, color=MUTED)


def add_inline(paragraph, text, size=11, color=BLACK):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=size - 0.5, color=DARK_BLUE)
            set_cell = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F6")
            set_cell.append(shd)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Callout"]
    p.paragraph_format.keep_together = True
    add_inline(p, text, size=12, color=INK)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    return p


def add_markdown_table(doc, rows):
    headers = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    data = [[c.strip() for c in row.strip().strip("|").split("|")] for row in rows[2:]]
    columns = len(headers)
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, value, size=9.5, color=INK)
        for run in p.runs:
            run.bold = True
    for row_values in data:
        row = table.add_row()
        for idx in range(columns):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, row_values[idx] if idx < len(row_values) else "", size=9.2)
    for cell in header.cells:
        set_cell_margins(cell, top=100, bottom=100)
    if columns == 4:
        widths = [1700, 2820, 2660, 2180]
    elif columns == 3:
        widths = [2600, 3180, 3580]
    elif columns == 2:
        widths = [2450, 6910]
    else:
        widths = [9360 // columns] * columns
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.keep_with_next = False
        style.paragraph_format.keep_together = False

    callout = doc.styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(12)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(12)
    callout.paragraph_format.line_spacing = 1.15

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8.5)
    code.font.color.rgb = INK
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("MARKET OPPORTUNITY PLATFORM  |  TRANSITION BRIEF")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "C8D0DA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("WORKING PROPOSAL  |  AUGUST 5, 2026  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)
    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ffp.add_run("WORKING PROPOSAL  |  HUMAN-REVIEWED DECISION SUPPORT")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STRATEGY AND PRODUCT HANDOFF")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("From Retail and Clinic Location Evaluator")
    set_run_font(run, size=26, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("to Market Opportunity Platform")
    set_run_font(run, size=26, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    run = p.add_run("Transition brief and new-project foundation")
    set_run_font(run, size=15, color=MUTED)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    values = [
        ("CURRENT FOUNDATION", "Clinic-first evidence, market comparison, deterministic scoring, and human review"),
        ("PROPOSED DIRECTION", "A broader platform for defining, comparing, and investigating market opportunities"),
    ]
    for row_idx, (label, value) in enumerate(values):
        left, right = table.rows[row_idx].cells
        set_cell_shading(left, "0B2545" if row_idx == 0 else "2E74B5")
        set_cell_shading(right, "F4F6F9")
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=130, bottom=130, start=140, end=140)
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=9, color=WHITE, bold=True)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        add_inline(rp, value, size=10.5, color=INK)
    set_table_geometry(table, [2200, 7160])
    set_table_borders(table, color="FFFFFF", size="8")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared August 5, 2026")
    set_run_font(run, size=10.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Working proposal for review")
    set_run_font(run, size=10.5, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    items = [
        "Executive summary",
        "1. Why the original project exists",
        "2. What has been learned",
        "3. Current project inventory",
        "4. What is demonstrated, proposed, and unresolved",
        "5. The proposed pivot",
        "6. Proposed platform workflow",
        "7. Proposed core domain model",
        "8. What to reuse, refactor, and leave behind",
        "9. Recommended MVP for the new project",
        "10. Measurement plan",
        "11. Governance and AI boundaries",
        "12. Suggested new-project structure",
        "13. New-project startup checklist",
        "14. Decisions required from the new project sponsor",
        "15. Copy-ready project charter",
        "Appendices A-C",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    add_inline(note, "This brief distinguishes current repository facts from proposed future-state choices. Source IDs refer to the repository source registry.", size=9.5, color=MUTED)
    doc.add_page_break()


def render_markdown(doc, lines):
    i = 0
    in_code = False
    code_lines = []
    active_num_id = None
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        if stripped.startswith("```"):
            active_num_id = None
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph(style="Code Block")
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F7")
                p_pr.append(shd)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Courier New", size=8.5, color=INK)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            active_num_id = None
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            active_num_id = None
            table_rows = [raw, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            active_num_id = None
            level = len(heading.group(1)) - 1
            text = heading.group(2)
            major_section = bool(re.match(r"^\d+\.\s", text)) or text.startswith("Appendix ")
            split_subsection = text in {
                "3.4 Technical foundation",
                "5.3 Provisional users and decisions",
                "Explicit non-goals for the first release",
                "First principle",
                "Source IDs to retain as historical context",
            }
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            if major_section or split_subsection:
                p.paragraph_format.page_break_before = True
            add_inline(p, text, size={1: 16, 2: 13, 3: 12}[min(level, 3)], color=BLUE if level < 3 else DARK_BLUE)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        if stripped.startswith("> "):
            active_num_id = None
            add_callout(doc, stripped[2:])
            i += 1
            continue

        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            add_inline(p, bullet.group(1))
            i += 1
            continue

        number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number:
            if active_num_id is None:
                active_num_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            apply_numbering(p, active_num_id)
            add_inline(p, number.group(1))
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "- ", "> ", "```", "|")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        text = " ".join(paragraph_lines)
        active_num_id = None
        p = doc.add_paragraph()
        add_inline(p, text)


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "From Retail and Clinic Location Evaluator to Market Opportunity Platform"
    props.subject = "Transition brief and new-project foundation"
    props.author = "Project team"
    props.keywords = "market opportunity, location evaluation, evidence, strategy, handoff"

    add_cover(doc)
    add_contents(doc)

    start = next(idx for idx, line in enumerate(lines) if line.strip() == "## Executive summary")
    render_markdown(doc, lines[start:])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
